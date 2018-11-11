#!/usr/bin/python
# -*- coding: cp936 -*-

import datetime
import os
import shutil
import sys
import time
import docx

def cpt(t):
	import json
	print(json.dumps(t,encoding="gbk",ensure_ascii=False))

class main:
	def __init__(self):
		if len(sys.argv)==1:
			self.arg=""
		else:
			self.arg=sys.argv[1]
		self.main()
	def main(self):
#		self.f=docx.Document("ID.docx")
		self.f=docx.Document("ofbdep.docx")
		print("�������:%d" %(len(self.f.tables)))
		for i in range(len(self.f.tables)):
			self.t=self.f.tables[i]
			if self.arg in ["","all"] or self.arg=="%d" %(i):
				self.printtable(i)
			if hasattr(main,"getinfo%d" %(i)):
				if self.arg=="%d" %(i) or self.arg=="all":
					print("::::::run getinfo%d" %(i))
					getattr(main,"getinfo%d" %(i))(self)
	def fmt(self,s):
		return s.replace("\n",",").encode("gbk").strip()
	def getfield(self,filename):	#ȡ���ļ����ֶα�
		srow=""
		f=open("../doc/%s" %(filename),"w")
		for j in range(len(self.t.rows)):
			r=self.t.rows[j]
			for cell in r.cells:
				t=cell.text.replace("\n",",").encode("gbk")
				srow=srow+t+"|"
			print(srow)
			f.write(srow+"\n")
			srow=""
		f.close()
	def getinfo4(self):	#����ҵ������ ����ҵ�����뼰ȷ�ϴ����
		srow=""
		f=open("../doc/servicetype.txt","w")
		for j in range(len(self.t.rows)):
			r=self.t.rows[j]
			for cell in r.cells:
				t=cell.text.replace("\n",",").encode("gbk")
				srow=srow+t+"|"
			print(srow)
			f.write(srow+"\n")
			srow=""
		f.close()
	def getinfo5(self):	#��������001���ʻ���Ϣ�޸�����003��������
		self.gettableinfo("st001.txt")
	def getinfo6(self):	#����ȷ��101��������
		self.gettableinfo("st101.txt")
	def getinfo7(self):	#��������002��������
		self.gettableinfo("st002.txt")
	def getinfo8(self):	#����ȷ��102��ҵ���������
		self.gettableinfo("st102.txt")
	def getinfo9(self):	#�˻���������004��������
		self.gettableinfo("st004.txt")
	def getinfo10(self):	#�˻��ⶳ����005 �˻����������007
		self.gettableinfo("st005.txt")
	def getinfo11(self):	#�����˻�����ʧ����006
		self.gettableinfo("st006.txt")
	def getinfo12(self):	#���ӽ����˻�����008
		self.gettableinfo("st008.txt")
	def getinfo13(self):	#��������˺�����058
		self.gettableinfo("st058.txt")
	def getinfo14(self):	#��������˺�ȷ��158
		self.gettableinfo("st158.txt")
	def getinfo15(self):	#�Ϲ�����(020)��ԤԼ�Ϲ�����(021)
		self.gettableinfo("st020.txt")
	def getinfo16(self):	#�Ϲ�ȷ��(021)��ԤԼ�Ϲ�ȷ��(121)
		self.gettableinfo("st021.txt")
	def getinfo17(self):	#�깺����(022)��ԤԼ�깺����(023)����ʱ�����깺����(039)
		self.gettableinfo("st022.txt")
	def getinfo18(self):	#�깺ȷ��(122)����ʱ�����깺ȷ��(139)
		srow=""
		f=open("../doc/%s" %("st122.txt"),"w")
		for j in range(len(self.t.rows)):
			r=self.t.rows[j]
			srow=self.fmt(r.cells[0].text)+"|"+self.fmt(r.cells[6].text)+"|"+self.fmt(r.cells[7].text)
			print(srow)
			f.write(srow+"\n")
			srow=""
		f.close()
	def getinfo19(self):	#ԤԼ�깺ȷ��123
		self.gettableinfo("st123.txt")
	def getinfo20(self):	#�������(024)����ʱ�����������(063)��ԤԼ�������(025)
		self.gettableinfo("st024.txt")
	def getinfo21(self):	#���ȷ��(124)����ʱ�������ȷ��(163)��ǿ�����ȷ��(142)
		self.gettableinfo("st124.txt")
	def getinfo22(self):	#ԤԼ���ȷ��(125)
		self.gettableinfo("st125.txt")
	def getinfo23(self):	#ת������/��������(026)��ת������/����ת������(027)��ת������/����ת������(028)
		self.gettableinfo("st026.txt")
	def getinfo24(self):	#ת������/����ȷ��(126)��ת������/����ת��ȷ��(127)��ת������/����ת��ȷ��(128)
		self.gettableinfo("st126.txt")
	def getinfo25(self):	#���÷ֺ췽ʽ����029
		self.gettableinfo("st029.txt")
	def getinfo26(self):	#���÷ֺ췽ʽȷ��129
		self.gettableinfo("st129.txt")
	def getinfo27(self):	#�Ϲ����130
		self.gettableinfo("st130.txt")
	def getinfo28(self):	#���������������031
		self.gettableinfo("st031.txt")
	def getinfo29(self):	#�����������ȷ��131
		self.gettableinfo("st131.txt")
	def getinfo30(self):	#��������ⶳ����032
		self.gettableinfo("st032.txt")
	def getinfo31(self):	#��������ⶳȷ��132
		self.gettableinfo("st132.txt")
	def getinfo32(self):	#�ǽ��׹�������(033)���ǽ��׹���ת������(034)���ǽ��׹���ת������(035)
		self.gettableinfo("st033.txt")
	def getinfo33(self):	#�ǽ��׹���ȷ��(133)���ǽ��׹���ת��ȷ��(134)���ǽ��׹���ת��ȷ��(135)
		self.gettableinfo("st133.txt")
	def getinfo34(self):	#����ת������(036)������ת��ת������(037)������ת��ת������(038)
		self.gettableinfo("st036.txt")
	def getinfo35(self):	#����ת��ȷ��(136)������ת��ת��ȷ��(137)������ת��ת��ȷ��(138)
		self.gettableinfo("st136.txt")
	def getinfo36(self):	#����/������Ͷ�ʷ���(143)
		self.gettableinfo("st143.txt")
	def getinfo37(self):	#ǿ�е���(144) ��ǿ�е���(145)
		self.gettableinfo("st144.txt")
	def getinfo38(self):	#���146
		self.gettableinfo("st146.txt")
	def getinfo39(self):	#ļ��ʧ��149
		self.gettableinfo("st149.txt")
	def getinfo40(self):	#��������(150)��������ֹ(151)
		self.gettableinfo("st150.txt")
	def getinfo41(self):	#����(052)����ԤԼ��(053)
		self.gettableinfo("st052.txt")
	def getinfo42(self):	#����ȷ��(152)����ԤԼ��ȷ��(153)
		self.gettableinfo("st152.txt")
	def getinfo43(self):	#��Ч�ʽ�(054)
		self.gettableinfo("st054.txt")
	def getinfo44(self):	#�����������ʽ����㣨155��
		self.gettableinfo("st155.txt")
	def getinfo45(self):	#Ͷ�����ʽ����㣨156��
		self.gettableinfo("st156.txt")
	def getinfo46(self):	#�����ⶳȷ�ϣ�157��
		self.gettableinfo("st157.txt")
	def getinfo47(self):	#��ʱ����ע�����루059������ʱ����ע�����루060��
		self.gettableinfo("st059.txt")
	def getinfo48(self):	#��ʱ����ע��ȷ�ϣ�159������ʱ����ע��ȷ�ϣ�160��
		self.gettableinfo("st159.txt")
	def getinfo49(self):	#��ʱ���������루061��
		self.gettableinfo("st061.txt")
	def getinfo50(self):	#��ʱ������ȷ�ϣ�161��
		self.gettableinfo("st161.txt")
	def getinfo51(self):	#�Ϲ���������(062)
		self.gettableinfo("st062.txt")
	def getinfo52(self):	#�Ϲ�����ȷ��(162)
		self.gettableinfo("st162.txt")
	def getinfo53(self):	#������������ͨ���루067���������������������루068��
		self.gettableinfo("st067.txt")
	def getinfo54(self):	#������������ͨȷ�ϣ�167������������������ȷ�ϣ�168��
		self.gettableinfo("st167.txt")
	def getinfo55(self):	#����ȷ�ϣ�169��
		self.gettableinfo("st169.txt")
	def getinfo69(self):	#01 �ʻ�����
		self.getfield("01.txt")
	def getinfo70(self):	#02 �ʻ�ȷ��
		self.getfield("02.txt")
	def getinfo71(self):	#03 ��������
		self.getfield("03.txt")
	def getinfo72(self):	#04 ����ȷ��
		self.getfield("04.txt")
	def getinfo73(self):	#05 ����
		self.getfield("05.txt")
	def getinfo74(self):	#06 �ֺ�
		self.getfield("06.txt")
	def getinfo75(self):	#07 ��̬��Ϣ�������顢��ֵ�ļ���
		self.getfield("07.txt")
	def getinfo76(self):	#08 ����
		self.getfield("08.txt")
	def getinfo77(self):	#09 ��������
		self.getfield("09.txt")
	def getinfo78(self):	#10 �ս�������
		self.getfield("10.txt")
	def getinfo79(self):	#11 TA���͵�ҵ���������
		self.getfield("11.txt")
	def getinfo80(self):	#12 TA���͵�ҵ��ȷ�ϻ���
		self.getfield("12.txt")
	def getinfo81(self):	#13 �����˷��͵�ҵ���������
		self.getfield("13.txt")
	def getinfo82(self):	#21 �����˼�����ϯλ
		self.getfield("21.txt")
	def getinfo83(self):	#23 ����������(����ȷȨ��������ҵ��
		self.getfield("23.txt")
	def getinfo84(self):	#24 ������ȷ��(����ȷȨ��������������ȷ��ҵ��
		self.getfield("24.txt")
	def getinfo85(self):	#25 �ʽ������ļ�
		self.getfield("25.txt")
	def getinfo86(self):	#26 ��Ʋ�Ʒ�ݶ���ϸ
		self.getfield("26.txt")
	def getinfo87(self):	#R1 �Ǿ�����˰��Ϣ�걨
		self.getfield("R1.txt")
	def getinfo88(self):	#R2 �Ǿ�����˰��Ϣȷ��
		self.getfield("R2.txt")
	def getinfo89(self):	#c1 �����������
		self.getfield("c1.txt")
	def getinfo90(self):	#c2 ��������ϵ
		self.getfield("c2.txt")
	def getinfo91(self):	#c3 ����ת����ϵ����
		self.getfield("c3.txt")
	def getinfo92(self):	#c4 ����ֺ췽������ ע��ȡ����c4�ļ�������
		self.getfield("c4.txt")
	def getinfo93(self):	#c5 �������
		self.getfield("c5.txt")
	def getinfo94(self):	#c6 ��Ʋ�Ʒ���������
		self.getfield("c6.txt")
	def getinfo95(self):	#ȡ�����ֵ��
		srow=""
		row=[]
		f=open("../doc/field.txt","w")
		for j in range(len(self.t.rows)):
			r=self.t.rows[j]
			for cell in r.cells:
				t=cell.text.replace("\n",",").encode("gbk").strip()
				if len(row)==1 and row[0]==t:	#�����word�ĵ���ʽ���ң���ʱ����һ�У�ֻ���������ı��취����
					continue
				if len(row)==2 and t not in ["A","C","N","����"]:
					print("��ʽ����:"+srow+t)
					sys.exit(1)
				if len(row)==3:
					num=filter(str.isdigit,t)
					row.append(num)
					srow=srow+num+"|"
					if row[2]=="N":
						ext=["��","һ","��","��","��","��","��","��","��","��"]
						for i in range(1,len(ext)):
							if t.find(ext[i])>0:
								row.append("%d" %(i))
								srow=srow+"%d" %(i)+"|"
								print("%s %d" %(t,i))
						if len(row)==4:
							row.append("0")
							srow=srow+"0|"
					else:
						row.append("")
						srow=srow+"|"
				else:
					row.append(t)
					srow=srow+t+"|"
				if len(row)==7:
					print(srow)
					f.write(srow+"\n")
					srow=""
					row=[]
	def gettableinfo(self,filename):	#ȡ�����ݣ����ڻ�ȡ���׶�Ӧ���ֶα�
		srow=""
		f=open("../doc/%s" %(filename),"w")
		for j in range(1,len(self.t.rows)):
			r=self.t.rows[j]
			srow=self.fmt(r.cells[0].text)+"|"+self.fmt(r.cells[5].text)+"|"+self.fmt(r.cells[6].text)
			print(srow)
			f.write(srow+"\n")
			srow=""
		f.close()
	def printtable(self,tno):	#������
		print("================��%d�ű�" %(tno))
		for j in range(len(self.t.rows)):
			if j>5:break
			r=self.t.rows[j]
			for cell in r.cells:
				print("%s|" %(cell.text)),
#			for k in range(len(r.cells)-2):
#				print("%d|%s|" %(k,r.cells[k].text))
#				print("%d|%s|" %(k,self.t.cell(j,k).text))
			print("")

main()
